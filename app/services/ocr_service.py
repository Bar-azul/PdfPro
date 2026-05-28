"""
OCRService — memory optimized + auto-rotation + image enhancement.
"""

import gc
import logging
import time
from pathlib import Path

import fitz
import pytesseract
from PIL import Image, ImageEnhance

from ..config import settings
from ..services.pdf_service import _temp_pdf, _temp_file, _ms

logger = logging.getLogger(__name__)
pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_PATH

IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp", ".webp", ".gif"}


def _is_image(path: Path) -> bool:
    return path.suffix.lower() in IMAGE_EXTS


def _image_to_pdf(image_path: Path) -> Path:
    with Image.open(image_path) as img:
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        elif img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        out = _temp_pdf("img_as_pdf")
        img.save(str(out), format="PDF", resolution=150)
    return out


def _prepare_image(img: Image.Image) -> Image.Image:
    """
    Prepare image for best OCR accuracy:
    1. Resize if too large
    2. Auto-detect and fix rotation
    3. Convert to grayscale
    4. Enhance contrast and sharpness
    """
    # Step 1 — Resize if too large
    max_px = 3000
    if max(img.width, img.height) > max_px:
        ratio = max_px / max(img.width, img.height)
        img = img.resize(
            (int(img.width * ratio), int(img.height * ratio)),
            Image.LANCZOS,
        )

    # Step 2 — Convert to RGB for OSD
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")

    # Step 3 — Auto-detect and fix rotation
    try:
        osd = pytesseract.image_to_osd(
            img,
            output_type=pytesseract.Output.DICT,
            config="--psm 0",
        )
        angle = osd.get("rotate", 0)
        if angle and angle != 0:
            img = img.rotate(-angle, expand=True)
            logger.info(f"Auto-rotated image by {angle}°")
    except Exception as e:
        logger.debug(f"OSD failed (continuing without rotation): {e}")

    # Step 4 — Grayscale + enhance
    img = img.convert("L")
    img = ImageEnhance.Contrast(img).enhance(2.0)
    img = ImageEnhance.Sharpness(img).enhance(2.0)

    return img


class OCRService:

    @staticmethod
    def extract_text(
        pdf_path: Path,
        language: str = "heb+eng",
        dpi: int = 200,
        pages: list[int] | None = None,
    ) -> list[dict]:
        t0 = time.time()

        # ── IMAGE: direct Tesseract — no PDF overhead ─────────────────────────
        if _is_image(pdf_path):
            result = OCRService.ocr_image(pdf_path, language=language)
            logger.info(f"Image OCR (direct) in {_ms(t0)}ms")
            return [{
                "page": 1,
                "text": result["text"],
                "confidence": result["confidence"],
                "source": "ocr",
            }]

        # ── PDF: page by page ─────────────────────────────────────────────────
        results = []
        matrix = fitz.Matrix(dpi / 72, dpi / 72)

        with fitz.open(pdf_path) as doc:
            target = [p - 1 for p in pages] if pages else range(doc.page_count)
            for i in target:
                if not (0 <= i < doc.page_count):
                    continue
                page = doc[i]

                # Use native text if available
                native_text = page.get_text().strip()
                if native_text and len(native_text) > 50:
                    results.append({
                        "page": i + 1,
                        "text": native_text,
                        "confidence": 1.0,
                        "source": "native",
                    })
                    continue

                # Render → prepare → OCR
                pix = page.get_pixmap(matrix=matrix, alpha=False)
                img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
                del pix

                img = _prepare_image(img)

                data = pytesseract.image_to_data(
                    img,
                    lang=language,
                    config="--oem 3 --psm 3",
                    output_type=pytesseract.Output.DICT,
                )
                del img

                words = [
                    w for w, c in zip(data["text"], data["conf"])
                    if w.strip() and int(c) > 20
                ]
                valid_confs = [int(c) for c in data["conf"] if int(c) > 0]
                avg_conf = (sum(valid_confs) / len(valid_confs) / 100) if valid_confs else 0.0

                results.append({
                    "page": i + 1,
                    "text": " ".join(words),
                    "confidence": round(avg_conf, 3),
                    "source": "ocr",
                })
                gc.collect()

        gc.collect()
        logger.info(f"OCR: {len(results)} pages, lang={language} in {_ms(t0)}ms")
        return results

    @staticmethod
    def extract_to_txt(
        pdf_path: Path, language: str = "heb+eng", dpi: int = 200
    ) -> Path:
        results = OCRService.extract_text(pdf_path, language=language, dpi=dpi)
        out = _temp_file("ocr_output", ".txt")
        lines = []
        for r in results:
            lines.append(f"=== עמוד {r['page']} ===")
            lines.append(r["text"])
            lines.append("")
        out.write_text("\n".join(lines), encoding="utf-8")
        gc.collect()
        return out

    @staticmethod
    def extract_to_searchable_pdf(
        pdf_path: Path, language: str = "heb+eng", dpi: int = 200
    ) -> Path:
        t0 = time.time()
        is_img = _is_image(pdf_path)

        # ── IMAGE: OCR directly → create PDF → add text layer ─────────────────
        if is_img:
            result = OCRService.ocr_image(pdf_path, language=language)
            ocr_text = result["text"]
            actual_path = _image_to_pdf(pdf_path)
            try:
                with fitz.open(actual_path) as doc:
                    if ocr_text.strip():
                        try:
                            doc[0].insert_text(
                                fitz.Point(10, 20),
                                ocr_text,
                                fontsize=1,
                                color=(1, 1, 1),
                                overlay=False,
                            )
                        except Exception as e:
                            logger.warning(f"insert_text failed: {e}")
                    out = _temp_pdf("searchable")
                    doc.save(out, deflate=True)
            finally:
                actual_path.unlink(missing_ok=True)
                gc.collect()
            logger.info(f"Searchable PDF (image) in {_ms(t0)}ms")
            return out

        # ── PDF: page by page OCR → add text layer ────────────────────────────
        results = OCRService.extract_text(pdf_path, language=language, dpi=dpi)
        text_by_page = {r["page"]: r["text"] for r in results}

        with fitz.open(pdf_path) as doc:
            for page_num, text in text_by_page.items():
                if not text.strip():
                    continue
                try:
                    doc[page_num - 1].insert_text(
                        fitz.Point(10, 20),
                        text,
                        fontsize=1,
                        color=(1, 1, 1),
                        overlay=False,
                    )
                except Exception as e:
                    logger.warning(f"insert_text failed page {page_num}: {e}")
            out = _temp_pdf("searchable")
            doc.save(out, deflate=True)

        gc.collect()
        logger.info(f"Searchable PDF in {_ms(t0)}ms")
        return out

    @staticmethod
    def extract_to_docx(
        pdf_path: Path, language: str = "heb+eng", dpi: int = 200
    ) -> Path:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        results = OCRService.extract_text(pdf_path, language=language, dpi=dpi)
        doc = Document()
        title = doc.add_heading("מסמך מחולץ — OCR", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        for r in results:
            doc.add_heading(f"עמוד {r['page']}", level=2)
            para = doc.add_paragraph(r["text"])
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in para.runs:
                run.font.name = "David"
                run.font.size = Pt(12)
            doc.add_paragraph()

        out = _temp_file("ocr_output", ".docx")
        doc.save(str(out))
        gc.collect()
        return out

    @staticmethod
    def ocr_image(image_path: Path, language: str = "heb+eng") -> dict:
        """
        Run Tesseract directly on an image file.
        Auto-detects rotation and enhances image quality before OCR.
        """
        t0 = time.time()

        with Image.open(image_path) as img:
            img = _prepare_image(img)
            data = pytesseract.image_to_data(
                img,
                lang=language,
                config="--oem 3 --psm 3",
                output_type=pytesseract.Output.DICT,
            )

        words = [
            w for w, c in zip(data["text"], data["conf"])
            if w.strip() and int(c) > 20
        ]
        text = " ".join(words)
        valid_confs = [int(c) for c in data["conf"] if int(c) > 0]
        avg_conf = (sum(valid_confs) / len(valid_confs) / 100) if valid_confs else 0.0

        gc.collect()
        logger.info(f"Image OCR in {_ms(t0)}ms, conf={avg_conf:.2f}")
        return {"text": text, "confidence": round(avg_conf, 3)}

    @staticmethod
    def get_available_languages() -> list[str]:
        try:
            langs = pytesseract.get_languages(config="")
            return [l for l in langs if l != "osd"]
        except Exception:
            return settings.OCR_LANGUAGES